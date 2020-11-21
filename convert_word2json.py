import json
import sys
import re
from docx import Document
from docx.shared import Inches
file=sys.argv[1]
output=sys.argv[2]

document = Document(file)
body_elements = document._body._body
#extract those wrapped in <w:r> tag
rs = body_elements.xpath('.//w:r')
#check if style is hyperlink (toc)
table_of_content = [r.text for r in rs if r.style == "Hyperlink"]

paragraphs=document.paragraphs
header_styles=[]
for i in range(0,len(paragraphs)):
	if document.paragraphs[i].style.name=='Title':
		header_styles.append(['Heading 0',document.paragraphs[i].text,0])
		original_heading=document.paragraphs[i].text

counter=1
for i in range(1,len(paragraphs)):
	if 'Heading' in document.paragraphs[i].style.name:
		header_styles.append([document.paragraphs[i].style.name,document.paragraphs[i].text,counter])
		counter+=1

texts=[]
counter=0
for i in range(1,len(paragraphs)):
	if 'Heading' in document.paragraphs[i].style.name or 'Title' in document.paragraphs[i].style.name:
		print(document.paragraphs[i].style.name)
		print(document.paragraphs[i].text)
		for j in header_styles:
			if document.paragraphs[i].text==j[1]:
				k=i+1
				paras=[]
				while k<len(paragraphs):
					print(document.paragraphs[k].text)
					if 'Heading' in document.paragraphs[k].style.name or 'Title' in document.paragraphs[k].style.name:
						paras.append(document.paragraphs[k].text)
						print('here')
						print(paras)
						if len(paras)>1:
							paras='\n'.join(paras)
						else:
							paras=paras[0]
						texts.append((j[1],paras))
						break
					else:
						print('there')
						paras.append(document.paragraphs[k].text)
					k+=1
		counter+=1

levels=[]
for i in range(0,len(header_styles)):
	header_styles[i][0]=int(header_styles[i][0].replace('Heading ',''))
	for j in texts:
		if header_styles[i][1]==j[0]:
			header_styles[i].append(j[1])
	levels.append(header_styles[i][0])


pairs=[]

def find_pairs(i_s,header_styles,levels):
	i_e=i_s
	while i_e > -1:
		level=1
		start,end = header_styles[i_s],header_styles[i_e]
		if end[0] < start[0]:
			print('match found')
			print(start,end)
			if len(start)==4:
				sentence=start[3]
			else:
				sentence=''
			return (start[2],end[2],levels[i_s]+1,start[1],sentence)
			break
		level+=1
		i_e-=1


last = len(header_styles)-1
pairs=[]
for i in range(last,0,-1):
	print(i)
	pairs.append(find_pairs(i,header_styles,levels))

pairs=pairs[::-1]
pairs.insert(0,(0,'NA',1,original_heading,''))
queue_tuples=pairs
pGraph = {"_id":0,"name":queue_tuples[0][3],"children":[]}

for tuple in queue_tuples:
	level = tuple[2]
	g = pGraph["children"]
	for i in range(level-1):
		for e in g:
			if tuple[1]==e["_id"]:
				g = e["children"]
	flag = 0
	for e in g:
		if tuple[0]==e["_id"]:
			flag = 1
			break
	if flag==0:
		if tuple[0]==0:
			_id='p1'
		else:
			_id=tuple[0]
		if tuple[1]==0:
			parent_id='p1'
		else:
			parent_id=tuple[1]
		if tuple[1]=='NA':
			g.append({"_id":_id,
				"name":tuple[3],
				'text':tuple[4],
				"children":[]})
		else:
			g.append({"_id":_id,
				"parentAreaRef":{'id':parent_id},
				"name":tuple[3],
				'text':tuple[4],
				"children":[]})


final=pGraph['children'][0]
for i in range(1,len(pGraph['children'])-1):
	final['children'].append(pGraph['children'][i])

with open(output, 'w') as f:
	json.dump(final,f,indent=4)


