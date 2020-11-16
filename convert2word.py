import json
import sys
import re
from docx import Document
from docx.shared import Inches
print(sys.argv)
file=sys.argv[1]
output=sys.argv[2]
print(file)

#file='/Users/Deirdreclarkson/js/json_files/GWAS2.YnP.json'
#output='GWAS2.docx'
with open(file,'r', encoding='utf-8',errors='ignore') as json_file:
	data = json.loads(json_file.read())

print(data)

document = Document()

document.add_heading(data['name'], 0)
document.add_paragraph(data['text'])

def add_header(header,counter):
	try:
		document.add_heading(header,level=counter)
	except ValueError:
		print('too deep')

def recursive_iter(obj,counter=0):
	if len(obj['children'])<1:
		print('end')
		add_header(obj['name'], counter)

		text=re.sub(r'\n','',obj['text'])
		document.add_paragraph(text)
		if 'linker' in obj:
			document.add_paragraph(obj['linker'])
	else:
		counter+=1
		print(counter)
		for i in obj['children']:
			print(i['name'])
			if len(i['children'])<1:
				add_header(i['name'], counter+1)
				# text=i['text']
				# print(text)
				if 'text' in i:
					text=i['text']
					print(text)
					if text!=None:
						text=i['text'].replace('\n','')
					document.add_paragraph(text)
				if 'linker' in i:
					document.add_paragraph(i['linker'])
				continue
			else:
				add_header(i['name'], counter)
				# text=i['text']
				# print(text)
				if 'text' in i:
					text=i['text']
					print(text)
					if text!=None:
						text=i['text'].replace('\n','')
					document.add_paragraph(text)
				if 'linker' in i:
					document.add_paragraph(i['linker'])
				recursive_iter(i,counter)

recursive_iter(data)
document.save(output)


